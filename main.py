from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Form, Request
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles  # ⭕ 1. 정적 파일 지원 보완
from sqlalchemy.orm import Session
from datetime import datetime, timedelta
from typing import Optional
import jwt
from pydantic import BaseModel
import os
from database import SessionLocal, User, Document
from llm_service import generate_summary
import io
import PyPDF2
import docx
import olefile

# --- CONFIG ---
SECRET_KEY = "docshub-local-secret-2025"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24  # 24시간

app = FastAPI(title="Document Hub API")

from fastapi.responses import FileResponse
import os

@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    return FileResponse(os.path.join(os.path.dirname(__file__), "favicon.png"))
    
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- DB DEPENDENCY ---
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="api/auth/login")

def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=15))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid token")
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")
        
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user

def get_optional_user(request: Request, db: Session = Depends(get_db)):
    """비로그인(게스트)과 로그인을 모두 수용하는 선택적 사용자 디펜던시"""
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        return None
    token = auth_header.split(" ")[1]
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if not username:
            return None
        return db.query(User).filter(User.username == username).first()
    except jwt.PyJWTError:
        return None

# --- PYDANTIC MODELS ---
class UserCreate(BaseModel):
    username: str
    password: str
    displayname: Optional[str] = ""
    operator_code: Optional[str] = None

class SummaryRequest(BaseModel):
    text: str
    prompt: str = ""
    model: Optional[str] = None

class QARequest(BaseModel):
    text: str
    question: str
    model: Optional[str] = None

class DocumentCreate(BaseModel):
    title: str
    original_content: str
    summary_content: str
    team_id: Optional[int] = None

class TeamCreate(BaseModel):
    name: str

class TeamJoinRequest(BaseModel):
    code: str

class PaymentRequest(BaseModel):
    plan_name: str
    token_amount: int

class ProfileUpdate(BaseModel):
    displayname: Optional[str] = None
    new_password: Optional[str] = None

class InquiryCreate(BaseModel):
    type: str
    title: str
    content: str

# ====================================================================
# AUTH ROUTES
# ====================================================================
@app.post("/api/auth/register", status_code=201)
def register(user: UserCreate, db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == user.username).first():
        raise HTTPException(status_code=400, detail="이미 사용 중인 아이디입니다.")
    is_operator = False
    if user.operator_code:
        secret_code = os.getenv("OPERATOR_CODE", "ADMIN1234")
        if user.operator_code == secret_code:
            is_operator = True
        else:
            raise HTTPException(status_code=400, detail="올바르지 않은 운영자 인증 코드입니다.")
            
    new_user = User(
        username=user.username,
        password_hash=user.password,  
        displayname=user.displayname or user.username,
        is_operator=is_operator
    )
    db.add(new_user)
    db.commit()
    return {"message": "회원가입 성공"}

@app.post("/api/auth/login")
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == form_data.username).first()
    if not user or user.password_hash != form_data.password:
        raise HTTPException(status_code=400, detail="아이디 또는 비밀번호가 올바르지 않습니다.")
    token = create_access_token(
        data={"sub": user.username},
        expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    )
    return {
        "access_token": token,
        "token_type": "bearer",
        "displayname": user.displayname or user.username,
        "tokens": user.tokens,
        "is_operator": getattr(user, "is_operator", False)
    }

@app.put("/api/auth/profile")
def update_profile(update: ProfileUpdate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    if update.displayname is not None:
        current_user.displayname = update.displayname
    if update.new_password:
        current_user.password_hash = update.new_password
    db.commit()
    return {"message": "프로필이 업데이트되었습니다.", "displayname": current_user.displayname}

@app.delete("/api/auth/account")
def delete_account(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """계정 삭제 API"""
    db.query(Document).filter(Document.owner_id == current_user.id).delete()
    from database import TeamMember
    db.query(TeamMember).filter(TeamMember.user_id == current_user.id).delete()
    db.delete(current_user)
    db.commit()
    return {"message": "계정이 삭제되었습니다."}

@app.get("/api/auth/me")
def get_me(current_user: User = Depends(get_current_user)):
    return {
        "id": current_user.id,
        "username": current_user.username,
        "displayname": current_user.displayname,
        "tokens": current_user.tokens,
        "summary_count": len(current_user.documents),
        "is_operator": getattr(current_user, "is_operator", False)
    }

# ====================================================================
# SUMMARY ROUTES (화면단에서 옛날 모델명을 강제로 보내도 서버에서 무조건 2.5 최신 모델로 고정합니다)
# ====================================================================
@app.post("/api/summary/generate")
def generate_summary_api(request: SummaryRequest, db: Session = Depends(get_db), current_user: Optional[User] = Depends(get_optional_user)):
    if current_user:
        if current_user.tokens <= 0:
            raise HTTPException(status_code=402, detail="보유 토큰이 소모되었습니다. 결제 페이지에서 충전해 주세요.")
        current_user.tokens -= 1
        db.commit()
        db.refresh(current_user)
    
    # [강제 조치] 화면에서 옛날 1.5 모델명을 보내오더라도 무조건 최신 gemini-2.5-flash로 강제 지정하여 404 에러를 방지합니다.
    actual_model = "gemini-2.5-flash"
    summary = generate_summary(request.text, request.prompt, actual_model)
    return {"summary": summary, "remaining_tokens": current_user.tokens if current_user else None}

@app.post("/api/summary/qa")
def document_qa_api(request: QARequest):
    """요약된 문서 기반 AI Q&A 피드백 엔드포인트"""
    from llm_service import answer_document_question
    # [강제 조치] 여기도 404 에러 방지를 위해 최신 모델로 강제 지정합니다.
    actual_model = "gemini-2.5-flash"
    answer = answer_document_question(request.text, request.question, actual_model)
    return {"answer": answer}

def extract_text_from_pdf(content_bytes: bytes) -> str:
    pdf_file = io.BytesIO(content_bytes)
    reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(content_bytes: bytes) -> str:
    docx_file = io.BytesIO(content_bytes)
    doc = docx.Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        if para.text:
            text += para.text + "\n"
    return text

def extract_text_from_hwp(content_bytes: bytes) -> str:
    hwp_file = io.BytesIO(content_bytes)
    ole = olefile.OleFileIO(hwp_file)
    if ole.exists('PrvText'):
        encoded_text = ole.openstream('PrvText').read()
        decoded_text = encoded_text.decode('utf-16')
        return decoded_text
    return ""

@app.post("/api/summary/upload")
async def upload_and_summarize(
    file: UploadFile = File(...),
    prompt: str = Form(""),
    model: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: Optional[User] = Depends(get_optional_user)
):
    if current_user:
        if current_user.tokens <= 0:
            raise HTTPException(status_code=402, detail="보유 토큰이 소모되었습니다. 결제 페이지에서 충전해 주세요.")
            
    filename = file.filename.lower()
    content = await file.read()
    text = ""
    try:
        if filename.endswith(".txt"):
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('cp949')
        elif filename.endswith(".pdf"):
            text = extract_text_from_pdf(content)
        elif filename.endswith(".docx"):
            text = extract_text_from_docx(content)
        elif filename.endswith(".hwp"):
            text = extract_text_from_hwp(content)
        else:
            raise HTTPException(status_code=400, detail="지원하지 않는 파일 형식입니다. (txt, pdf, docx, hwp 지원)")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"파일에서 텍스트를 추출하는 중 오류가 발생했습니다: {str(e)}")
        
    if not text.strip():
        raise HTTPException(status_code=400, detail="파일에서 텍스트를 추출할 수 없습니다. 빈 파일이거나 암호화된 문서일 수 있습니다.")
        
    if current_user:
        current_user.tokens -= 1
        db.commit()
        db.refresh(current_user)
        
    # [강제 조치] 파일 업로드 요약 시에도 구버전 404 에러를 차단합니다.
    actual_model = "gemini-2.5-flash"
    summary = generate_summary(text, prompt, actual_model)
    return {"original_text": text, "summary": summary, "remaining_tokens": current_user.tokens if current_user else None}

# ====================================================================
# DOCUMENT ROUTES
# ====================================================================
@app.post("/api/documents", status_code=201)
def save_document(doc: DocumentCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """FR-03: 요약본 DB 저장 (개인 또는 팀 문서)"""
    new_doc = Document(
        title=doc.title,
        original_content=doc.original_content,
        summary_content=doc.summary_content,
        owner_id=current_user.id,
        team_id=doc.team_id
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)
    return new_doc

@app.get("/api/documents/my")
def get_my_documents(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """내 요약 문서 목록 (개인 문서)"""
    docs = db.query(Document).filter(Document.owner_id == current_user.id, Document.team_id == None).order_by(Document.created_at.desc()).all()
    return [{
        "id": d.id,
        "title": d.title,
        "summary_content": d.summary_content,
        "is_shared": d.is_shared,
        "created_at": d.created_at.isoformat(),
        "owner_id": d.owner_id
    } for d in docs]

@app.get("/api/documents/{doc_id}")
def get_document(doc_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id, Document.owner_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")
    return doc

@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """문서 삭제 API"""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")
    if doc.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="문서를 삭제할 권한이 없습니다.")
    db.delete(doc)
    db.commit()
    return {"message": "문서가 성공적으로 삭제되었습니다."}

@app.put("/api/documents/{doc_id}/share")
def toggle_share(doc_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """FR-04, FR-06: 공유 설정"""
    doc = db.query(Document).filter(Document.id == doc_id, Document.owner_id == current_user.id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="문서를 찾을 수 없습니다.")
    doc.is_shared = not doc.is_shared
    db.commit()
    return {"is_shared": doc.is_shared}

@app.get("/api/documents/shared")
def get_shared_documents(query: str = "", db: Session = Depends(get_db)):
    """FR-06: 공유된 요약본 검색"""
    q = db.query(Document).filter(Document.is_shared == True)
    if query:
        q = q.filter(Document.title.contains(query))
    return q.order_by(Document.created_at.desc()).all()

# ====================================================================
# TEAM COLABORATION ROUTES
# ====================================================================
import random
import string

def generate_team_code():
    return "T-" + "".join(random.choices(string.ascii_uppercase + string.digits, k=6))

@app.post("/api/teams", status_code=201)
def create_team(team_req: TeamCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """새로운 팀 생성 및 생성자를 첫 멤버로 추가"""
    from database import Team, TeamMember
    attempts = 0
    code = generate_team_code()
    while db.query(Team).filter(Team.code == code).first() and attempts < 10:
        code = generate_team_code()
        attempts += 1
        
    new_team = Team(name=team_req.name, code=code, owner_id=current_user.id)
    db.add(new_team)
    db.commit()
    db.refresh(new_team)
    
    member = TeamMember(team_id=new_team.id, user_id=current_user.id)
    db.add(member)
    db.commit()
    return new_team

@app.post("/api/teams/join")
def join_team(join_req: TeamJoinRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """초대 코드를 사용하여 팀 가입"""
    from database import Team, TeamMember
    team = db.query(Team).filter(Team.code == join_req.code).first()
    if not team:
        raise HTTPException(status_code=404, detail="해당 코드를 가진 팀을 찾을 수 없습니다.")
        
    existing = db.query(TeamMember).filter(TeamMember.team_id == team.id, TeamMember.user_id == current_user.id).first()
    if existing:
        return {"message": "이미 이 팀의 멤버입니다.", "team_id": team.id, "team_name": team.name}
        
    new_member = TeamMember(team_id=team.id, user_id=current_user.id)
    db.add(new_member)
    db.commit()
    return {"message": f"'{team.name}' 팀에 성공적으로 가입되었습니다.", "team_id": team.id, "team_name": team.name}

@app.delete("/api/teams/{team_id}/leave")
def leave_team(team_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """팀 탈퇴 또는 해체 API"""
    from database import Team, TeamMember
    team = db.query(Team).filter(Team.id == team_id).first()
    if not team:
        raise HTTPException(status_code=404, detail="팀을 찾을 수 없습니다.")
    member = db.query(TeamMember).filter(TeamMember.team_id == team_id, TeamMember.user_id == current_user.id).first()
    if not member:
        raise HTTPException(status_code=400, detail="팀의 멤버가 아닙니다.")
        
    if team.owner_id == current_user.id:
        db.delete(team)
        db.commit()
        return {"message": "팀 소유자이므로 팀이 해체되었습니다.", "disbanded": True}
    else:
        db.delete(member)
        db.commit()
        return {"message": "팀에서 탈퇴했습니다.", "disbanded": False}

@app.get("/api/teams")
def list_my_teams(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """내가 소속된 모든 팀 목록 조회"""
    from database import Team, TeamMember
    memberships = db.query(TeamMember).filter(TeamMember.user_id == current_user.id).all()
    team_ids = [m.team_id for m in memberships]
    teams = db.query(Team).filter(Team.id.in_(team_ids)).order_by(Team.created_at.desc()).all()
    result = []
    for t in teams:
        m_count = db.query(TeamMember).filter(TeamMember.team_id == t.id).count()
        result.append({
            "id": t.id,
            "name": t.name,
            "code": t.code,
            "owner_id": t.owner_id,
            "created_at": t.created_at.isoformat(),
            "member_count": m_count
        })
    return result

@app.get("/api/teams/{team_id}/documents")
def get_team_documents(team_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """특정 팀에 소속된 모든 요약 문서 목록 조회"""
    from database import TeamMember, Document
    member = db.query(TeamMember).filter(TeamMember.team_id == team_id, TeamMember.user_id == current_user.id).first()
    if not member:
        raise HTTPException(status_code=403, detail="이 팀의 문서를 볼 권한이 없습니다.")
    docs = db.query(Document).filter(Document.team_id == team_id).order_by(Document.created_at.desc()).all()
    return [{
        "id": d.id,
        "title": d.title,
        "summary_content": d.summary_content,
        "is_shared": d.is_shared,
        "created_at": d.created_at.isoformat(),
        "owner_name": d.owner.displayname or d.owner.username,
        "owner_id": d.owner_id
    } for d in docs]

@app.get("/api/teams/{team_id}/members")
def get_team_members(team_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """특정 팀의 모든 멤버 목록 조회"""
    from database import TeamMember
    member = db.query(TeamMember).filter(TeamMember.team_id == team_id, TeamMember.user_id == current_user.id).first()
    if not member:
        raise HTTPException(status_code=403, detail="이 팀의 멤버 목록을 볼 권한이 없습니다.")
    members = db.query(TeamMember).filter(TeamMember.team_id == team_id).all()
    return [{
        "username": m.user.username,
        "displayname": m.user.displayname or m.user.username,
        "joined_at": m.joined_at.isoformat()
    } for m in members]

# ====================================================================
# INQUIRY ROUTES
# ====================================================================
@app.post("/api/inquiries", status_code=201)
def create_inquiry(inquiry: InquiryCreate, db: Session = Depends(get_db), current_user: Optional[User] = Depends(get_optional_user)):
    from database import Inquiry
    new_inquiry = Inquiry(
        type=inquiry.type,
        title=inquiry.title,
        content=inquiry.content,
        user_id=current_user.id if current_user else None
    )
    db.add(new_inquiry)
    db.commit()
    db.refresh(new_inquiry)
    return {"message": "문의가 성공적으로 접수되었습니다."}

@app.get("/api/inquiries")
def get_inquiries(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    if not getattr(current_user, "is_operator", False):
        raise HTTPException(status_code=403, detail="운영자만 접근 가능합니다.")
    from database import Inquiry
    inquiries = db.query(Inquiry).order_by(Inquiry.created_at.desc()).all()
    return [{
        "id": i.id,
        "type": i.type,
        "title": i.title,
        "content": i.content,
        "created_at": i.created_at.isoformat(),
        "username": i.user.username if i.user else "비회원",
        "displayname": (i.user.displayname or i.user.username) if i.user else "비회원"
    } for i in inquiries]

# ====================================================================
# MOCK PAYMENT ROUTES
# ====================================================================
@app.post("/api/payment/charge")
def charge_tokens(pay_req: PaymentRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """가상 결제 승인 후 토큰 충전"""
    current_user.tokens += pay_req.token_amount
    db.commit()
    db.refresh(current_user)
    return {
        "message": f"'{pay_req.plan_name}' 결제가 완료되어 {pay_req.token_amount} 토큰이 충전되었습니다.",
        "tokens": current_user.tokens
    }

# ----------------------------------------------------------------7
# ✨ [신규 추가]: 구글 OAuth 2.0 소셜 로그인 API 라우터
# ----------------------------------------------------------------
import httpx
from pydantic import BaseModel

